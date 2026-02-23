import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
import re
import os
import zlib
import base64

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.append(start)
    text = OxmlElement('w:t')
    text.text = bookmark_text
    run._r.append(text)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    end.set(qn('w:name'), bookmark_name)
    tag.append(end)

def create_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def enable_update_fields_on_open(doc):
    # This setting forces Microsoft Word to update the Table of Contents upon opening of the file
    element = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    element.append(update_fields)

def encode_kroki(text):
    compressed = zlib.compress(text.encode('utf-8'))
    return base64.urlsafe_b64encode(compressed).decode('utf-8')

def parse_markdown_to_docx(filepath, doc):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_lang = ''
    code_lines = []
    
    for line in lines:
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                if code_lang == 'mermaid':
                    mermaid_code = '\n'.join(code_lines)
                    try:
                        encoded = encode_kroki(mermaid_code)
                        url = f"https://kroki.io/mermaid/png/{encoded}"
                        
                        import urllib.request
                        temp_img = "temp_mermaid.png"
                        
                        req = urllib.request.Request(
                            url, 
                            data=None, 
                            headers={
                                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                            }
                        )
                        with urllib.request.urlopen(req) as response, open(temp_img, 'wb') as out_file:
                            data = response.read()
                            out_file.write(data)
                            
                        # Add white background to make lines visible
                        try:
                            from PIL import Image
                            img = Image.open(temp_img)
                            if img.mode in ('RGBA', 'LA'):
                                bg = Image.new('RGB', img.size, (255, 255, 255))
                                bg.paste(img, mask=img.split()[-1])
                                bg.save(temp_img)
                        except Exception as img_e:
                            print(f"Failed to process image background: {img_e}")
                        
                        doc.add_picture(temp_img, width=Inches(6.0))
                        
                        if os.path.exists(temp_img):
                            os.remove(temp_img)
                    except Exception as e:
                        print(f"Failed to render mermaid diagram: {e}")
                        doc.add_paragraph(f"Failed to render mermaid diagram: {e}")
                else:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    p.paragraph_format.left_indent = Inches(0.5)
                    for code_line in code_lines:
                        p.add_run(code_line + '\n').font.name = 'Courier New'
            continue
            
        if in_code_block:
             code_lines.append(line)
             continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
             # Basic handling for bold text in list items
             p = doc.add_paragraph(style='List Bullet')
             text = line[2:].strip()
             parts = re.split(r'(\*\*.*?\*\*)', text)
             for part in parts:
                 if part.startswith('**') and part.endswith('**'):
                     p.add_run(part[2:-2]).bold = True
                 else:
                     p.add_run(part)
        elif line.strip() == '':
            doc.add_paragraph()
        else:
            # Handle inline bold in normal paragraphs
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)


def build_solution_arch_doc():
    doc = docx.Document()
    
    # Title Page
    title = doc.add_heading('Solution Architecture Document', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('\n\n\n\n')
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    create_toc(doc)
    doc.add_page_break()

    # Summary Section
    doc.add_heading('Summary', level=1)
    p = doc.add_paragraph("This document provides a comprehensive overview of the ChatApp solution architecture. It combines the project's README, functional and security requirements, and deployment requirements into a single consolidated reference.")
    doc.add_page_break()

    # Parse Files
    files = ['../README.md', '../requirements.md', '../DeploymentRequirements.md']
    for file in files:
        if os.path.exists(file):
            parse_markdown_to_docx(file, doc)
            doc.add_page_break()
            
    enable_update_fields_on_open(doc)
    
    doc.save('../SolutionArchitecture.docx')
    print("Successfully generated SolutionArchitecture.docx in the project root")

if __name__ == "__main__":
    build_solution_arch_doc()
